from docx import Document
import configparser
import os
import sys
import console

version={
    "release":0,
    "major":0,
    "minor":0,
    "patch":0
}

defaults=configparser.ConfigParser()
defaults.sections()
defaults.read('defaults.conf')

document=Document()

def __init__(inputFile):
    if os.path.exists('TMP/'):
        if os.path.exists('TMP/output'+defaults['EXTENSIONS']['MicrosoftWord']):
            os.remove('TMP/output'+defaults['EXTENSIONS']['MicrosoftWord'])
    else:
        os.mkdir('TMP')

    file=open(inputFile,'r')
    count = 0
    while True:
        count += 1
        line = file.readline()
        if not line:
            file.close()
            document.save('TMP/output'+defaults['EXTENSIONS']['MicrosoftWord'])
            break

        # <head>
        # <style>
        if line.strip().startswith('DH'):
            console.error('Unsupported element at line: '+str(count))
        elif line.strip().startswith('DW'):
            console.error('Unsupported element at line: '+str(count))
        elif line.strip().startswith('BG'):
            console.error('Unsupported element at line: '+str(count))
        elif line.strip().startswith('TXT'):
            console.error('Unsupported element at line: '+str(count))
        # </style>
        elif line.strip().startswith('T'):
            console.error('Unsupported element at line: '+str(count))
        elif line.strip().startswith('ICO'):
            console.error('Unsupported element at line: '+str(count))
        # </head>

        elif line.strip().startswith('H1'):
            document.add_heading(line.strip().split(':;:')[1],level=1)
            console.log('Wrote "Hedding 1" element')
        elif line.strip().startswith('H2'):
            document.add_heading(line.strip().split(':;:')[1],level=2)
            console.log('Wrote "Hedding 2" element')
        elif line.strip().startswith('H3'):
            document.add_heading(line.strip().split(':;:')[1],level=3)
            console.log('Wrote "Hedding 3" element')
        elif line.strip().startswith('P'):
            document.add_paragraph(line.strip().split(':;:')[1])
            console.log('Wrote "Paragraph" element')
        elif line.strip().startswith('C'):
            console.warn('Comment at line '+str(count))
        elif line.strip().startswith('~'):
            console.info('line '+str(count)+' is blank')
            document.add_paragraph()
        else:
            print()
            sys.exit('Line '+str(count)+" contains an invalid element, or it's blank")